"""
Analyse des documents joints (PDF, Excel, Word) via Claude.

Jusqu'ici les documents étaient téléchargés mais jamais lus : un bon de
pesée, une facture ou une fiche de déclassement était classé sur son seul
nom de fichier. Ce service lit le CONTENU :

  - PDF  : envoyé directement à Claude (gère le texte ET le scanné, pas
           besoin d'OCR séparé).
  - XLSX : texte extrait avec openpyxl, puis structuré par Claude.
  - DOCX : texte extrait avec python-docx, puis structuré par Claude.

Sortie structurée (type de doc, résumé, réf, client, site, matière,
quantité, montant, anomalie) stockée dans `document_analysis`, puis
fusionnée dans la classification du message.
"""

from __future__ import annotations

import base64
import json
import logging
from io import BytesIO

import anthropic

from app.config import settings
from app.db import get_supabase

logger = logging.getLogger(__name__)

DOCUMENT_MODEL = "claude-sonnet-4-5"

# Claude accepte jusqu'à ~32 MB / 100 pages pour un PDF en base64.
_MAX_PDF_BYTES = 30 * 1024 * 1024
# Au-delà, on tronque le texte Office envoyé au modèle.
_MAX_OFFICE_CHARS = 24_000


SYSTEM_PROMPT = """Tu es un analyseur de documents pour une entreprise française de \
collecte/évacuation de déchets en Île-de-France (PVC, ferraille, alu, bennes grutables). \
Les documents sont des bons de pesée, bons de livraison, bordereaux de suivi de déchets \
(BSD), fiches de déclassement, factures, bons de commande.

Ton rôle : transformer le document en JSON structuré exploitable. Tu lis ce qui est \
ÉCRIT, factuellement, sans inventer.

Règles strictes :
1. Tu réponds UNIQUEMENT par un objet JSON conforme au schéma. Aucun texte avant/après.
2. Si une information n'apparaît pas, mets null. Jamais d'invention.
3. `summary` : 1-2 phrases factuelles en français, < 250 caractères.
4. `full_text` : le texte intégral lisible du document (transcris fidèlement). Tronque \
   si très long mais garde les infos clés (réf, quantités, parties, dates).
5. `possible_anomaly` : true si quelque chose cloche (document illisible, incomplet, \
   signature/cachet manquant, incohérence de quantité, doublon évident).
6. `confidence` : 0.0 (incertain) à 1.0 (sûr).

Schéma JSON :
{
  "document_type": "bon_pesee | bon_livraison | fiche_declassement | facture | bsd | bon_commande | autre",
  "summary": "<phrase factuelle>",
  "reference": "<n° BL / réf / n° facture ou null>",
  "client_name": "<client ou null>",
  "site_name": "<chantier/site ou null>",
  "waste_type": "<matière/déchet ou null>",
  "quantity": "<quantité avec unité ou null>",
  "amount": "<montant ou null>",
  "doc_dates": ["<dates mentionnées>"],
  "full_text": "<texte intégral lisible>",
  "possible_anomaly": true/false,
  "anomaly_description": "<si anomalie ou null>",
  "confidence": 0.0-1.0
}"""


def _is_pdf(mime_type: str, filename: str | None) -> bool:
    if "pdf" in (mime_type or "").lower():
        return True
    return bool(filename and filename.lower().endswith(".pdf"))


def _is_xlsx(mime_type: str, filename: str | None) -> bool:
    mt = (mime_type or "").lower()
    if "spreadsheet" in mt or "excel" in mt:
        return True
    return bool(filename and filename.lower().endswith((".xlsx", ".xlsm")))


def _is_docx(mime_type: str, filename: str | None) -> bool:
    mt = (mime_type or "").lower()
    if "wordprocessing" in mt or "msword" in mt:
        return True
    return bool(filename and filename.lower().endswith(".docx"))


def _extract_xlsx_text(file_bytes: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"# Feuille : {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
            if sum(len(x) for x in lines) > _MAX_OFFICE_CHARS:
                lines.append("… (tronqué)")
                return "\n".join(lines)
    return "\n".join(lines)


def _extract_docx_text(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(file_bytes))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return text[:_MAX_OFFICE_CHARS]


def _parse_json(raw: str) -> dict:
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.strip("`")
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()
    return json.loads(raw)


async def _call_claude(content_blocks: list[dict]) -> dict | None:
    client = anthropic.AsyncAnthropic(api_key=settings.anthropic_api_key)
    try:
        response = await client.messages.create(
            model=DOCUMENT_MODEL,
            max_tokens=2048,
            system=[
                {
                    "type": "text",
                    "text": SYSTEM_PROMPT,
                    "cache_control": {"type": "ephemeral"},
                }
            ],
            messages=[{"role": "user", "content": content_blocks}],
        )
        return _parse_json(response.content[0].text)
    except json.JSONDecodeError as exc:
        logger.warning("Document: JSON invalide renvoyé par Claude: %s", exc)
        return None
    except Exception as exc:  # noqa: BLE001
        logger.exception("Document: appel Claude échoué: %s", exc)
        return None
    finally:
        await client.close()


async def analyze_document(
    media_id: str,
    file_bytes: bytes,
    mime_type: str,
    filename: str | None = None,
) -> dict | None:
    """
    Analyse un document joint, stocke le résultat dans document_analysis,
    et renvoie un dict (avec kind='document') pour fusion dans la classification.
    Best-effort : renvoie None sans lever en cas d'échec.
    """
    if not settings.anthropic_api_key:
        logger.warning("ANTHROPIC_API_KEY absente, skip analyse document")
        return None
    if not file_bytes or len(file_bytes) < 50:
        return None

    result: dict | None = None
    try:
        if _is_pdf(mime_type, filename):
            if len(file_bytes) > _MAX_PDF_BYTES:
                logger.warning("PDF trop volumineux (%d B), skip", len(file_bytes))
                return None
            b64 = base64.standard_b64encode(file_bytes).decode("ascii")
            result = await _call_claude([
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": b64,
                    },
                },
                {"type": "text", "text": "Analyse ce document et produis le JSON."},
            ])
        elif _is_xlsx(mime_type, filename) or _is_docx(mime_type, filename):
            text = (
                _extract_xlsx_text(file_bytes)
                if _is_xlsx(mime_type, filename)
                else _extract_docx_text(file_bytes)
            )
            if not text.strip():
                return None
            result = await _call_claude([
                {
                    "type": "text",
                    "text": (
                        f"Nom du fichier : {filename or '(inconnu)'}\n\n"
                        f"Contenu extrait :\n\n{text}\n\nProduis le JSON."
                    ),
                },
            ])
        else:
            # Format non géré (txt, csv, etc.) — pas d'analyse mais pas d'erreur
            logger.info("Document type non géré pour analyse: %s / %s", mime_type, filename)
            return None
    except Exception as exc:  # noqa: BLE001
        logger.exception("analyze_document a échoué pour %s: %s", media_id, exc)
        return None

    if not result:
        return None

    row = {
        "media_id": media_id,
        "document_type": result.get("document_type"),
        "summary": result.get("summary"),
        "reference": result.get("reference"),
        "client_name": result.get("client_name"),
        "site_name": result.get("site_name"),
        "waste_type": result.get("waste_type"),
        "quantity": result.get("quantity"),
        "amount": result.get("amount"),
        "doc_dates": result.get("doc_dates") or [],
        "full_text": (result.get("full_text") or "")[:8000] or None,
        "possible_anomaly": bool(result.get("possible_anomaly")),
        "anomaly_description": result.get("anomaly_description"),
        "confidence": float(result.get("confidence", 0.0)),
        "model_used": DOCUMENT_MODEL,
    }

    try:
        sb = get_supabase()
        sb.table("document_analysis").upsert(row, on_conflict="media_id").execute()
    except Exception as exc:  # noqa: BLE001
        logger.exception("Stockage document_analysis échoué pour %s: %s", media_id, exc)

    logger.info(
        "Document analysé %s: type=%s anomaly=%s conf=%.2f",
        media_id, row["document_type"], row["possible_anomaly"], row["confidence"],
    )
    return {**row, "kind": "document"}
